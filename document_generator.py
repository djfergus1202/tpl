import io
import zipfile
from typing import Dict, Any, List, Optional, Union
from datetime import datetime
import re
import json
from dataclasses import dataclass
import os
import tempfile

# For Word document generation
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

@dataclass
class DocumentMetadata:
    """Document metadata container"""
    title: str
    author: str
    subject: str
    keywords: List[str]
    created_date: datetime
    document_type: str
    citation_style: str
    word_count: int

@dataclass
class DocumentSection:
    """Document section container"""
    title: str
    content: str
    level: int
    include_in_toc: bool = True
    page_break: bool = False

class DocumentGenerator:
    """Generate academic documents in various formats"""
    
    def __init__(self):
        self.default_styles = self._load_default_styles()
        self.templates = self._load_templates()
        self.citation_styles = self._load_citation_styles()
    
    def _load_default_styles(self) -> Dict[str, Dict[str, Any]]:
        """Load default document styles"""
        return {
            'academic': {
                'font_name': 'Times New Roman',
                'font_size': 12,
                'line_spacing': 1.5,
                'margins': {'top': 1.0, 'bottom': 1.0, 'left': 1.0, 'right': 1.0},
                'paragraph_spacing': {'before': 6, 'after': 6}
            },
            'thesis': {
                'font_name': 'Times New Roman',
                'font_size': 12,
                'line_spacing': 2.0,
                'margins': {'top': 1.0, 'bottom': 1.0, 'left': 1.5, 'right': 1.0},
                'paragraph_spacing': {'before': 0, 'after': 12}
            },
            'journal': {
                'font_name': 'Arial',
                'font_size': 11,
                'line_spacing': 1.15,
                'margins': {'top': 1.0, 'bottom': 1.0, 'left': 1.0, 'right': 1.0},
                'paragraph_spacing': {'before': 3, 'after': 3}
            }
        }
    
    def _load_templates(self) -> Dict[str, Dict[str, Any]]:
        """Load document templates"""
        return {
            'research_article': {
                'sections': [
                    'Title Page',
                    'Abstract',
                    'Keywords',
                    'Introduction',
                    'Methods',
                    'Results',
                    'Discussion',
                    'Conclusion',
                    'References',
                    'Appendices'
                ],
                'required_sections': ['Title Page', 'Abstract', 'Introduction', 'Methods', 'Results', 'Discussion', 'References'],
                'max_abstract_words': 250,
                'citation_style': 'APA'
            },
            'thesis_chapter': {
                'sections': [
                    'Chapter Title',
                    'Introduction',
                    'Literature Review',
                    'Methodology',
                    'Results and Analysis',
                    'Discussion',
                    'Conclusion',
                    'References'
                ],
                'required_sections': ['Chapter Title', 'Introduction', 'Methodology', 'Conclusion'],
                'citation_style': 'APA'
            },
            'grant_proposal': {
                'sections': [
                    'Title Page',
                    'Executive Summary',
                    'Problem Statement',
                    'Literature Review',
                    'Objectives',
                    'Methodology',
                    'Timeline',
                    'Budget',
                    'References',
                    'Appendices'
                ],
                'required_sections': ['Title Page', 'Executive Summary', 'Problem Statement', 'Objectives', 'Methodology', 'Budget'],
                'citation_style': 'APA'
            },
            'conference_abstract': {
                'sections': [
                    'Title',
                    'Background',
                    'Methods',
                    'Results',
                    'Conclusions'
                ],
                'required_sections': ['Title', 'Background', 'Methods', 'Results', 'Conclusions'],
                'max_words': 300,
                'citation_style': 'Vancouver'
            },
            'systematic_review': {
                'sections': [
                    'Title Page',
                    'Abstract',
                    'Introduction',
                    'Methods',
                    'Results',
                    'Discussion',
                    'Conclusion',
                    'References',
                    'PRISMA Checklist',
                    'Appendices'
                ],
                'required_sections': ['Title Page', 'Abstract', 'Introduction', 'Methods', 'Results', 'Discussion', 'References'],
                'citation_style': 'APA'
            }
        }
    
    def _load_citation_styles(self) -> Dict[str, Dict[str, str]]:
        """Load citation style formats"""
        return {
            'APA': {
                'in_text': '(Author, Year)',
                'reference_format': 'Author, A. A. (Year). Title. Journal Name, Volume(Issue), pages.',
                'book_format': 'Author, A. A. (Year). Title. Publisher.',
                'website_format': 'Author, A. A. (Year, Month Date). Title. Website Name. URL'
            },
            'MLA': {
                'in_text': '(Author Page)',
                'reference_format': 'Author, First. "Title." Journal Name, vol. #, no. #, Year, pp. #-#.',
                'book_format': 'Author, First. Title. Publisher, Year.',
                'website_format': 'Author, First. "Title." Website Name, Date, URL.'
            },
            'Chicago': {
                'in_text': '(Author Year)',
                'reference_format': 'Author, First Last. "Title." Journal Name vol, no. # (Year): pages.',
                'book_format': 'Author, First Last. Title. City: Publisher, Year.',
                'website_format': 'Author, First Last. "Title." Website Name. Date. URL.'
            },
            'Vancouver': {
                'in_text': '[#]',
                'reference_format': '#. Author AA. Title. Journal Name. Year;Vol(Issue):pages.',
                'book_format': '#. Author AA. Title. City: Publisher; Year.',
                'website_format': '#. Author AA. Title [Internet]. Website; Year [cited Year Month Date]. Available from: URL'
            },
            'Nature': {
                'in_text': '#',
                'reference_format': '#. Author, A. A. Title. Journal Name vol, pages (year).',
                'book_format': '#. Author, A. A. Title (Publisher, year).',
                'website_format': '#. Author, A. A. Title. Website Name URL (year).'
            }
        }
    
    def create_word_document(self, content: str, metadata: Optional[DocumentMetadata] = None,
                           template_type: str = 'academic', style_name: str = 'academic') -> bytes:
        """Create a Word document from content"""
        
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx package is required for Word document generation")
        
        # Create document
        doc = Document()
        
        # Set default metadata
        if metadata:
            doc.core_properties.title = metadata.title
            doc.core_properties.author = metadata.author
            doc.core_properties.subject = metadata.subject
            doc.core_properties.keywords = '; '.join(metadata.keywords)
            doc.core_properties.created = metadata.created_date
        
        # Apply styles
        self._apply_word_styles(doc, style_name)
        
        # Add title page if metadata provided
        if metadata and metadata.title:
            self._add_title_page(doc, metadata)
        
        # Process content
        if isinstance(content, str):
            sections = self._parse_content_sections(content)
        else:
            sections = content
        
        # Add sections
        for section in sections:
            self._add_word_section(doc, section)
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes.getvalue()
    
    def _apply_word_styles(self, doc: 'Document', style_name: str):
        """Apply styles to Word document"""
        
        if style_name not in self.default_styles:
            style_name = 'academic'
        
        style_config = self.default_styles[style_name]
        
        # Set default font for Normal style
        style = doc.styles['Normal']
        font = style.font
        font.name = style_config['font_name']
        font.size = Pt(style_config['font_size'])
        
        # Set paragraph formatting
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = style_config['line_spacing']
        paragraph_format.space_before = Pt(style_config['paragraph_spacing']['before'])
        paragraph_format.space_after = Pt(style_config['paragraph_spacing']['after'])
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(style_config['margins']['top'])
            section.bottom_margin = Inches(style_config['margins']['bottom'])
            section.left_margin = Inches(style_config['margins']['left'])
            section.right_margin = Inches(style_config['margins']['right'])
        
        # Create heading styles
        for i in range(1, 4):
            try:
                heading_style = doc.styles[f'Heading {i}']
                heading_font = heading_style.font
                heading_font.name = style_config['font_name']
                heading_font.size = Pt(style_config['font_size'] + (4 - i) * 2)
                heading_font.bold = True
            except:
                pass  # Style might not exist
    
    def _add_title_page(self, doc: 'Document', metadata: DocumentMetadata):
        """Add title page to Word document"""
        
        # Title
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(metadata.title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        
        # Author
        doc.add_paragraph()  # Spacing
        author_paragraph = doc.add_paragraph()
        author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_paragraph.add_run(metadata.author)
        author_run.font.size = Pt(14)
        
        # Date
        doc.add_paragraph()  # Spacing
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_paragraph.add_run(metadata.created_date.strftime('%B %d, %Y'))
        date_run.font.size = Pt(12)
        
        # Page break
        doc.add_page_break()
    
    def _parse_content_sections(self, content: str) -> List[DocumentSection]:
        """Parse content into sections"""
        
        sections = []
        
        # Split content by headers
        lines = content.split('\n')
        current_section = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            
            # Check if line is a header
            if self._is_header_line(line):
                # Save previous section
                if current_section:
                    sections.append(DocumentSection(
                        title=current_section,
                        content='\n'.join(current_content),
                        level=self._get_header_level(current_section)
                    ))
                
                # Start new section
                current_section = line
                current_content = []
            else:
                current_content.append(line)
        
        # Add final section
        if current_section:
            sections.append(DocumentSection(
                title=current_section,
                content='\n'.join(current_content),
                level=self._get_header_level(current_section)
            ))
        
        return sections
    
    def _is_header_line(self, line: str) -> bool:
        """Check if line is a header"""
        
        # Common header patterns
        header_patterns = [
            r'^#{1,6}\s+',  # Markdown headers
            r'^\d+\.\s+[A-Z]',  # Numbered headers
            r'^[A-Z][A-Z\s]+$',  # ALL CAPS headers
            r'^[A-Z][a-z\s]+:$',  # Title case with colon
        ]
        
        for pattern in header_patterns:
            if re.match(pattern, line):
                return True
        
        # Check for common section titles
        section_titles = [
            'abstract', 'introduction', 'background', 'methods', 'methodology',
            'results', 'findings', 'discussion', 'conclusion', 'references',
            'bibliography', 'appendix', 'acknowledgments'
        ]
        
        return line.lower().strip() in section_titles
    
    def _get_header_level(self, header: str) -> int:
        """Get header level (1-6)"""
        
        # Markdown style
        if header.startswith('#'):
            return min(6, header.count('#'))
        
        # Numbered headers
        if re.match(r'^\d+\.', header):
            return 1
        
        # Default level
        return 1
    
    def _add_word_section(self, doc: 'Document', section: DocumentSection):
        """Add section to Word document"""
        
        # Add heading
        if section.title:
            heading = doc.add_heading(section.title, level=section.level)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add content paragraphs
        paragraphs = section.content.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                paragraph = doc.add_paragraph(para_text)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Add page break if requested
        if section.page_break:
            doc.add_page_break()
    
    def create_latex_document(self, content: str, metadata: Optional[DocumentMetadata] = None,
                            template_type: str = 'article', citation_style: str = 'APA') -> str:
        """Create LaTeX document from content"""
        
        # Document class and packages
        latex_content = self._get_latex_preamble(template_type, citation_style)
        
        # Document metadata
        if metadata:
            latex_content += f"\\title{{{metadata.title}}}\n"
            latex_content += f"\\author{{{metadata.author}}}\n"
            latex_content += f"\\date{{{metadata.created_date.strftime('%B %d, %Y')}}}\n"
        
        # Begin document
        latex_content += "\n\\begin{document}\n\n"
        
        # Title page
        if metadata and metadata.title:
            latex_content += "\\maketitle\n\n"
        
        # Process content
        sections = self._parse_content_sections(content)
        
        for section in sections:
            latex_content += self._section_to_latex(section)
        
        # Bibliography
        latex_content += "\n\\bibliographystyle{plain}\n"
        latex_content += "\\bibliography{references}\n\n"
        
        # End document
        latex_content += "\\end{document}\n"
        
        return latex_content
    
    def _get_latex_preamble(self, template_type: str, citation_style: str) -> str:
        """Get LaTeX preamble"""
        
        document_classes = {
            'article': 'article',
            'thesis': 'report',
            'book': 'book',
            'conference': 'article'
        }
        
        doc_class = document_classes.get(template_type, 'article')
        
        preamble = f"\\documentclass[12pt,a4paper]{{{doc_class}}}\n"
        preamble += "\\usepackage[utf8]{inputenc}\n"
        preamble += "\\usepackage[T1]{fontenc}\n"
        preamble += "\\usepackage{amsmath,amsfonts,amssymb}\n"
        preamble += "\\usepackage{graphicx}\n"
        preamble += "\\usepackage{hyperref}\n"
        preamble += "\\usepackage{geometry}\n"
        preamble += "\\usepackage{setspace}\n"
        
        # Citation packages based on style
        if citation_style.upper() == 'APA':
            preamble += "\\usepackage{apacite}\n"
        elif citation_style.upper() == 'MLA':
            preamble += "\\usepackage{mla}\n"
        else:
            preamble += "\\usepackage{natbib}\n"
        
        # Page geometry
        preamble += "\\geometry{margin=1in}\n"
        
        # Line spacing
        if template_type == 'thesis':
            preamble += "\\doublespacing\n"
        else:
            preamble += "\\onehalfspacing\n"
        
        return preamble + "\n"
    
    def _section_to_latex(self, section: DocumentSection) -> str:
        """Convert section to LaTeX"""
        
        latex = ""
        
        # Section header
        if section.title:
            if section.level == 1:
                latex += f"\\section{{{section.title}}}\n\n"
            elif section.level == 2:
                latex += f"\\subsection{{{section.title}}}\n\n"
            elif section.level == 3:
                latex += f"\\subsubsection{{{section.title}}}\n\n"
            else:
                latex += f"\\paragraph{{{section.title}}}\n\n"
        
        # Section content
        content = self._escape_latex(section.content)
        paragraphs = content.split('\n\n')
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if paragraph:
                latex += f"{paragraph}\n\n"
        
        return latex
    
    def _escape_latex(self, text: str) -> str:
        """Escape special LaTeX characters"""
        
        escape_chars = {
            '&': '\\&',
            '%': '\\%',
            '$': '\\$',
            '#': '\\#',
            '^': '\\textasciicircum{}',
            '_': '\\_',
            '{': '\\{',
            '}': '\\}',
            '~': '\\textasciitilde{}',
            '\\': '\\textbackslash{}'
        }
        
        for char, escape in escape_chars.items():
            text = text.replace(char, escape)
        
        return text
    
    def create_html_document(self, content: str, metadata: Optional[DocumentMetadata] = None,
                           template_type: str = 'academic') -> str:
        """Create HTML document from content"""
        
        html = self._get_html_template(template_type)
        
        # Insert metadata
        if metadata:
            html = html.replace('{{TITLE}}', metadata.title)
            html = html.replace('{{AUTHOR}}', metadata.author)
            html = html.replace('{{DATE}}', metadata.created_date.strftime('%B %d, %Y'))
        else:
            html = html.replace('{{TITLE}}', 'Document')
            html = html.replace('{{AUTHOR}}', '')
            html = html.replace('{{DATE}}', '')
        
        # Process content
        sections = self._parse_content_sections(content)
        
        content_html = ""
        for section in sections:
            content_html += self._section_to_html(section)
        
        html = html.replace('{{CONTENT}}', content_html)
        
        return html
    
    def _get_html_template(self, template_type: str) -> str:
        """Get HTML template"""
        
        return """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{TITLE}}</title>
    <style>
        body {
            font-family: 'Times New Roman', serif;
            font-size: 12pt;
            line-height: 1.5;
            max-width: 8.5in;
            margin: 0 auto;
            padding: 1in;
            background-color: white;
            color: black;
        }
        
        .title-page {
            text-align: center;
            margin-bottom: 2in;
        }
        
        .title {
            font-size: 16pt;
            font-weight: bold;
            margin-bottom: 0.5in;
        }
        
        .author {
            font-size: 14pt;
            margin-bottom: 0.5in;
        }
        
        .date {
            font-size: 12pt;
        }
        
        h1 {
            font-size: 14pt;
            font-weight: bold;
            margin: 1em 0 0.5em 0;
            page-break-after: avoid;
        }
        
        h2 {
            font-size: 13pt;
            font-weight: bold;
            margin: 0.8em 0 0.4em 0;
            page-break-after: avoid;
        }
        
        h3 {
            font-size: 12pt;
            font-weight: bold;
            margin: 0.6em 0 0.3em 0;
            page-break-after: avoid;
        }
        
        p {
            margin: 0 0 0.5em 0;
            text-align: justify;
            text-indent: 0;
        }
        
        .abstract {
            margin: 1em 0;
            padding: 1em;
            background-color: #f9f9f9;
            border-left: 4px solid #ccc;
        }
        
        .references {
            margin-top: 2em;
        }
        
        .reference-item {
            margin-bottom: 1em;
            padding-left: 1em;
            text-indent: -1em;
        }
        
        @media print {
            body {
                margin: 0;
                padding: 1in;
            }
            
            .page-break {
                page-break-before: always;
            }
        }
    </style>
</head>
<body>
    <div class="title-page">
        <div class="title">{{TITLE}}</div>
        <div class="author">{{AUTHOR}}</div>
        <div class="date">{{DATE}}</div>
    </div>
    
    <div class="content">
        {{CONTENT}}
    </div>
</body>
</html>"""
    
    def _section_to_html(self, section: DocumentSection) -> str:
        """Convert section to HTML"""
        
        html = ""
        
        # Section header
        if section.title:
            header_tag = f"h{min(6, section.level)}"
            html += f"<{header_tag}>{section.title}</{header_tag}>\n"
        
        # Section content
        paragraphs = section.content.split('\n\n')
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if paragraph:
                # Simple formatting
                paragraph = self._apply_html_formatting(paragraph)
                html += f"<p>{paragraph}</p>\n"
        
        return html
    
    def _apply_html_formatting(self, text: str) -> str:
        """Apply basic HTML formatting"""
        
        # Bold text
        text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)
        text = re.sub(r'__(.*?)__', r'<strong>\1</strong>', text)
        
        # Italic text
        text = re.sub(r'\*(.*?)\*', r'<em>\1</em>', text)
        text = re.sub(r'_(.*?)_', r'<em>\1</em>', text)
        
        # Escape HTML characters
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        
        return text
    
    def create_markdown_document(self, content: str, metadata: Optional[DocumentMetadata] = None) -> str:
        """Create Markdown document from content"""
        
        markdown = ""
        
        # Front matter
        if metadata:
            markdown += "---\n"
            markdown += f"title: \"{metadata.title}\"\n"
            markdown += f"author: \"{metadata.author}\"\n"
            markdown += f"date: \"{metadata.created_date.strftime('%Y-%m-%d')}\"\n"
            if metadata.keywords:
                keywords_str = ', '.join(f'"{k}"' for k in metadata.keywords)
                markdown += f"keywords: [{keywords_str}]\n"
            markdown += "---\n\n"
        
        # Add content
        markdown += content
        
        return markdown
    
    def create_document_package(self, content: str, metadata: Optional[DocumentMetadata] = None,
                              formats: List[str] = None, template_type: str = 'academic') -> bytes:
        """Create a package with multiple document formats"""
        
        if formats is None:
            formats = ['docx', 'pdf', 'html', 'latex', 'markdown']
        
        # Create zip file in memory
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            
            # Generate each format
            for format_type in formats:
                try:
                    if format_type.lower() == 'docx' and DOCX_AVAILABLE:
                        doc_content = self.create_word_document(content, metadata, template_type)
                        zip_file.writestr('document.docx', doc_content)
                    
                    elif format_type.lower() == 'html':
                        html_content = self.create_html_document(content, metadata, template_type)
                        zip_file.writestr('document.html', html_content.encode('utf-8'))
                    
                    elif format_type.lower() == 'latex':
                        latex_content = self.create_latex_document(content, metadata, template_type)
                        zip_file.writestr('document.tex', latex_content.encode('utf-8'))
                    
                    elif format_type.lower() == 'markdown':
                        md_content = self.create_markdown_document(content, metadata)
                        zip_file.writestr('document.md', md_content.encode('utf-8'))
                    
                    elif format_type.lower() == 'txt':
                        # Plain text version
                        plain_text = self._content_to_plain_text(content)
                        zip_file.writestr('document.txt', plain_text.encode('utf-8'))
                
                except Exception as e:
                    # Create error log for failed formats
                    error_msg = f"Error generating {format_type}: {str(e)}"
                    zip_file.writestr(f'error_{format_type}.log', error_msg.encode('utf-8'))
            
            # Add metadata file
            if metadata:
                metadata_json = {
                    'title': metadata.title,
                    'author': metadata.author,
                    'subject': metadata.subject,
                    'keywords': metadata.keywords,
                    'created_date': metadata.created_date.isoformat(),
                    'document_type': metadata.document_type,
                    'citation_style': metadata.citation_style,
                    'word_count': metadata.word_count
                }
                zip_file.writestr('metadata.json', json.dumps(metadata_json, indent=2).encode('utf-8'))
        
        zip_buffer.seek(0)
        return zip_buffer.getvalue()
    
    def _content_to_plain_text(self, content: str) -> str:
        """Convert content to plain text"""
        
        # Remove markdown formatting
        text = re.sub(r'#{1,6}\s+', '', content)  # Headers
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
        text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
        text = re.sub(r'__(.*?)__', r'\1', text)  # Bold
        text = re.sub(r'_(.*?)_', r'\1', text)  # Italic
        
        return text
    
    def apply_template(self, content: str, template_name: str) -> Dict[str, str]:
        """Apply a document template to content"""
        
        if template_name not in self.templates:
            return {'error': f'Template {template_name} not found'}
        
        template = self.templates[template_name]
        sections = self._parse_content_sections(content)
        
        # Map content to template sections
        template_content = {}
        
        for template_section in template['sections']:
            # Find matching section in content
            matching_section = None
            for section in sections:
                if self._sections_match(section.title, template_section):
                    matching_section = section
                    break
            
            if matching_section:
                template_content[template_section] = matching_section.content
            else:
                template_content[template_section] = f"[{template_section} content to be added]"
        
        # Check for required sections
        missing_required = []
        for required in template.get('required_sections', []):
            if required not in template_content or not template_content[required].strip():
                missing_required.append(required)
        
        result = {
            'template_name': template_name,
            'sections': template_content,
            'missing_required': missing_required,
            'citation_style': template.get('citation_style', 'APA')
        }
        
        # Add template-specific validations
        if template_name == 'conference_abstract':
            word_count = sum(len(content.split()) for content in template_content.values())
            max_words = template.get('max_words', 300)
            if word_count > max_words:
                result['warnings'] = [f'Abstract exceeds maximum word count ({word_count}/{max_words} words)']
        
        return result
    
    def _sections_match(self, content_title: str, template_section: str) -> bool:
        """Check if content section matches template section"""
        
        content_lower = content_title.lower().strip()
        template_lower = template_section.lower().strip()
        
        # Direct match
        if content_lower == template_lower:
            return True
        
        # Fuzzy matching for common variations
        variations = {
            'introduction': ['intro', 'background'],
            'methodology': ['methods', 'method'],
            'results': ['findings', 'results and analysis'],
            'discussion': ['discussion and conclusion'],
            'conclusion': ['conclusions', 'summary'],
            'references': ['bibliography', 'works cited'],
            'abstract': ['summary', 'executive summary']
        }
        
        for standard, variants in variations.items():
            if template_lower == standard and content_lower in variants:
                return True
            if content_lower == standard and template_lower in variants:
                return True
        
        return False
    
    def validate_document_structure(self, content: str, template_name: str = None) -> Dict[str, Any]:
        """Validate document structure"""
        
        sections = self._parse_content_sections(content)
        
        validation_result = {
            'total_sections': len(sections),
            'word_count': len(content.split()),
            'issues': [],
            'suggestions': []
        }
        
        # Check for basic structure
        if len(sections) < 3:
            validation_result['issues'].append("Document appears to have very few sections")
        
        # Check for common academic sections
        section_titles = [s.title.lower() for s in sections]
        common_sections = ['introduction', 'methods', 'results', 'discussion', 'conclusion']
        
        missing_common = [sec for sec in common_sections if sec not in section_titles]
        if missing_common:
            validation_result['suggestions'].append(f"Consider adding these common sections: {', '.join(missing_common)}")
        
        # Template-specific validation
        if template_name and template_name in self.templates:
            template = self.templates[template_name]
            required_sections = template.get('required_sections', [])
            
            missing_required = []
            for required in required_sections:
                if not any(self._sections_match(s.title, required) for s in sections):
                    missing_required.append(required)
            
            if missing_required:
                validation_result['issues'].append(f"Missing required sections for {template_name}: {', '.join(missing_required)}")
        
        return validation_result
    
    def generate_table_of_contents(self, content: str) -> str:
        """Generate table of contents"""
        
        sections = self._parse_content_sections(content)
        
        toc = "Table of Contents\n"
        toc += "=" * 17 + "\n\n"
        
        page_num = 1
        
        for section in sections:
            if section.include_in_toc and section.title:
                indent = "  " * (section.level - 1)
                toc += f"{indent}{section.title} ... {page_num}\n"
                
                # Estimate page length (rough calculation)
                words_per_page = 250
                section_words = len(section.content.split())
                section_pages = max(1, section_words // words_per_page)
                page_num += section_pages
        
        return toc
    
    def count_words(self, content: str) -> Dict[str, int]:
        """Count words in document sections"""
        
        sections = self._parse_content_sections(content)
        
        word_counts = {
            'total': len(content.split()),
            'sections': {}
        }
        
        for section in sections:
            if section.title:
                word_counts['sections'][section.title] = len(section.content.split())
        
        return word_counts
