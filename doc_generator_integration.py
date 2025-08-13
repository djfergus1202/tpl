
import streamlit as st
import io
from datetime import datetime
try:
    from .utils.document_generator import DocumentGenerator
    from .utils.citation_validator import CitationValidator
except Exception:
    DocumentGenerator = None
    CitationValidator = None

def main():
    st.header("📄 Paper & Report Generator")
    st.write("Generate DOCX, PDF-ready LaTeX, and HTML reports with citations.")

    title = st.text_input("Title", value="Untitled Research Report")
    author = st.text_input("Author", value="Researcher")
    abstract = st.text_area("Abstract", height=150)
    body = st.text_area("Body (Markdown/Plain Text)", height=300)
    export_fmt = st.selectbox("Export Format", ["DOCX", "HTML", "Markdown"])

    if st.button("Generate Document", type="primary"):
        if DocumentGenerator is None:
            st.error("Document generator module not available.")
            return

        dg = DocumentGenerator()
        # Minimal example using markdown/HTML or docx
        if export_fmt == "Markdown":
            content = f"# {title}\n\n**{author}** — {datetime.now().date()}\n\n## Abstract\n{abstract}\n\n## Body\n{body}"
            st.download_button("Download .md", content, file_name="paper.md")
        elif export_fmt == "HTML":
            html = f"""<html><head><meta charset='utf-8'><title>{title}</title></head>
            <body><h1>{title}</h1><p><em>{author}</em> — {datetime.now().date()}</p>
            <h2>Abstract</h2><p>{abstract}</p><h2>Body</h2><pre>{body}</pre></body></html>"""
            st.download_button("Download .html", html, file_name="paper.html")
        else:
            # DOCX (best-effort; if python-docx missing, show message)
            try:
                from docx import Document
                doc = Document()
                doc.add_heading(title, 0)
                doc.add_paragraph(f"{author} — {datetime.now().date()}" )
                doc.add_heading("Abstract", level=1)
                doc.add_paragraph(abstract)
                doc.add_heading("Body", level=1)
                doc.add_paragraph(body)
                buf = io.BytesIO()
                doc.save(buf); buf.seek(0)
                st.download_button("Download .docx", buf.getvalue(), file_name="paper.docx")
            except Exception as e:
                st.error("python-docx not installed. Select HTML/Markdown, or add python-docx to requirements.")
                st.exception(e)
